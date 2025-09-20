import os
import re
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from bs4 import BeautifulSoup
import markdown
from utils import (
    clean_text, get_file_hash, ensure_directory, 
    sanitize_filename, format_file_size, validate_query
)

logger = logging.getLogger(__name__)

class DocumentLoader:
    """
    Loads documents from a local directory and splits them into chunks
    for embedding and retrieval.
    """

    def __init__(self, config: Dict[str, Any]):
        """
        Initialize DocumentLoader with configuration.
        
        :param config: Configuration dictionary containing ingestion settings
        """
        self.config = config
        self.data_dir = config.get('data_dir', 'data')
        self.chunk_size = config.get('chunk_size', 1000)
        self.chunk_overlap = config.get('chunk_overlap', 200)
        self.supported_formats = config.get('supported_formats', ['.pdf', '.docx', '.txt', '.md'])
        self.max_file_size_mb = config.get('max_file_size_mb', 50)
        self.enable_metadata_extraction = config.get('enable_metadata_extraction', True)
        self.preserve_formatting = config.get('preserve_formatting', True)
        
        # Ensure data directory exists
        ensure_directory(self.data_dir)
        
        logger.info(f"DocumentLoader initialized with data_dir: {self.data_dir}")
        logger.info(f"Supported formats: {self.supported_formats}")
        logger.info(f"Chunk size: {self.chunk_size}, Overlap: {self.chunk_overlap}")

    def _read_pdf(self, path: str) -> str:
        """Extract text from PDF file with enhanced error handling."""
        text = ""
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num} in {path}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error reading PDF file {path}: {e}")
            raise
        return clean_text(text)

    def _read_docx(self, path: str) -> str:
        """Extract text from DOCX file with enhanced formatting preservation."""
        text = ""
        try:
            doc = docx.Document(path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        
        except Exception as e:
            logger.error(f"Error reading DOCX file {path}: {e}")
            raise
        return clean_text(text)

    def _read_txt(self, path: str) -> str:
        """Read plain text file with encoding detection."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(path, "r", encoding=encoding) as f:
                        return clean_text(f.read())
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file {path} with any supported encoding")
        except Exception as e:
            logger.error(f"Error reading text file {path}: {e}")
            raise

    def _read_markdown(self, path: str) -> str:
        """Read and parse Markdown file."""
        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            
            if self.preserve_formatting:
                # Convert markdown to HTML then extract text
                html = markdown.markdown(content)
                soup = BeautifulSoup(html, 'html.parser')
                return clean_text(soup.get_text())
            else:
                return clean_text(content)
        except Exception as e:
            logger.error(f"Error reading Markdown file {path}: {e}")
            raise

    def _read_html(self, path: str) -> str:
        """Extract text from HTML file."""
        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return clean_text(text)
        except Exception as e:
            logger.error(f"Error reading HTML file {path}: {e}")
            raise

    def _extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file."""
        if not self.enable_metadata_extraction:
            return {}
        
        try:
            stat = os.stat(file_path)
            file_size = stat.st_size
            modified_time = stat.st_mtime
            
            metadata = {
                'file_size': file_size,
                'file_size_formatted': format_file_size(file_size),
                'modified_time': modified_time,
                'file_hash': get_file_hash(file_path),
                'file_extension': Path(file_path).suffix.lower()
            }
            
            # Extract additional metadata based on file type
            if file_path.lower().endswith('.pdf'):
                try:
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        if reader.metadata:
                            metadata.update({
                                'title': reader.metadata.get('/Title', ''),
                                'author': reader.metadata.get('/Author', ''),
                                'subject': reader.metadata.get('/Subject', ''),
                                'creator': reader.metadata.get('/Creator', ''),
                                'producer': reader.metadata.get('/Producer', ''),
                                'creation_date': str(reader.metadata.get('/CreationDate', '')),
                                'modification_date': str(reader.metadata.get('/ModDate', ''))
                            })
                        metadata['page_count'] = len(reader.pages)
                except Exception as e:
                    logger.warning(f"Could not extract PDF metadata from {file_path}: {e}")
            
            return metadata
        except Exception as e:
            logger.warning(f"Could not extract metadata from {file_path}: {e}")
            return {}

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks with intelligent boundary detection.
        """
        if not text.strip():
            return []
        
        # Use sentence-aware chunking if possible
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # If adding this sentence would exceed chunk size, start a new chunk
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # Fallback to simple chunking if sentence-based chunking fails
        if not chunks:
            chunks = []
            start = 0
            while start < len(text):
                end = min(start + self.chunk_size, len(text))
                chunks.append(text[start:end].strip())
                start += self.chunk_size - self.chunk_overlap
        
        return [chunk for chunk in chunks if chunk.strip()]

    def load(self) -> List[Dict]:
        """
        Load all documents from data_dir, split into chunks, and return a list of dicts.
        
        Returns:
            List of document chunks with structure:
            {
                'id': str,
                'text': str,
                'meta': {
                    'source': str,
                    'chunk_index': int,
                    'total_chunks': int,
                    'file_metadata': dict,
                    'keywords': list
                }
            }
        """
        docs = []
        processed_files = 0
        total_chunks = 0
        
        logger.info(f"Starting document ingestion from {self.data_dir}")
        
        # Get all files in the directory
        try:
            files = [f for f in os.listdir(self.data_dir) if os.path.isfile(os.path.join(self.data_dir, f))]
        except Exception as e:
            logger.error(f"Error reading directory {self.data_dir}: {e}")
            return []
        
        for fname in files:
            file_path = os.path.join(self.data_dir, fname)
            
            # Check file size
            try:
                file_size = os.path.getsize(file_path)
                if file_size > self.max_file_size_mb * 1024 * 1024:
                    logger.warning(f"Skipping {fname}: file too large ({format_file_size(file_size)})")
                    continue
            except Exception as e:
                logger.warning(f"Could not check size of {fname}: {e}")
                continue
            
            # Check if file format is supported
            file_ext = Path(fname).suffix.lower()
            if file_ext not in self.supported_formats:
                logger.debug(f"Skipping {fname}: unsupported format {file_ext}")
                continue
            
            try:
                # Read file content based on extension
                text = self._read_file_by_extension(file_path, file_ext)
                
                if not text or not text.strip():
                    logger.warning(f"Skipping {fname}: no text content extracted")
                    continue
                
                # Extract metadata
                metadata = self._extract_metadata(file_path)
                
                # Chunk the text
                chunks = self._chunk_text(text)
                
                if not chunks:
                    logger.warning(f"Skipping {fname}: no valid chunks created")
                    continue
                
                # Create document objects for each chunk
                for i, chunk in enumerate(chunks):
                    doc_id = f"{sanitize_filename(fname)}_chunk{i}"
                    
                    # Extract keywords from chunk
                    keywords = self._extract_keywords_from_chunk(chunk)
                    
                    doc = {
                        "id": doc_id,
                        "text": chunk,
                        "meta": {
                            "source": fname,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "file_metadata": metadata,
                            "keywords": keywords,
                            "chunk_length": len(chunk),
                            "file_path": file_path
                        }
                    }
                    docs.append(doc)
                    total_chunks += 1
                
                processed_files += 1
                logger.info(f"Processed {fname}: {len(chunks)} chunks created")
                
            except Exception as e:
                logger.error(f"Error processing file {fname}: {e}")
                continue
        
        logger.info(f"Document ingestion complete: {processed_files} files processed, {total_chunks} chunks created")
        return docs

    def _read_file_by_extension(self, file_path: str, file_ext: str) -> str:
        """Read file content based on file extension."""
        if file_ext == '.pdf':
            return self._read_pdf(file_path)
        elif file_ext == '.docx':
            return self._read_docx(file_path)
        elif file_ext == '.txt':
            return self._read_txt(file_path)
        elif file_ext == '.md':
            return self._read_markdown(file_path)
        elif file_ext == '.html':
            return self._read_html(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")

    def _extract_keywords_from_chunk(self, text: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from a text chunk."""
        try:
            # Simple keyword extraction using word frequency
            words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
            
            # Common stop words
            stop_words = {
                'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'its', 'may', 'new', 'now', 'old', 'see', 'two', 'who', 'boy', 'did', 'man', 'men', 'put', 'say', 'she', 'too', 'use'
            }
            
            # Count word frequencies
            word_freq = {}
            for word in words:
                if word not in stop_words and len(word) > 2:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Return top keywords
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            return [word for word, freq in sorted_words[:max_keywords]]
        except Exception as e:
            logger.warning(f"Error extracting keywords: {e}")
            return []

    def get_document_stats(self, docs: List[Dict]) -> Dict[str, Any]:
        """Get statistics about the loaded documents."""
        if not docs:
            return {}
        
        total_chunks = len(docs)
        total_text_length = sum(len(doc['text']) for doc in docs)
        
        # Count by file type
        file_types = {}
        sources = set()
        
        for doc in docs:
            source = doc['meta']['source']
            sources.add(source)
            file_ext = doc['meta']['file_metadata'].get('file_extension', 'unknown')
            file_types[file_ext] = file_types.get(file_ext, 0) + 1
        
        return {
            'total_documents': len(sources),
            'total_chunks': total_chunks,
            'total_text_length': total_text_length,
            'average_chunk_length': total_text_length / total_chunks if total_chunks > 0 else 0,
            'file_types': file_types,
            'sources': list(sources)
        }
